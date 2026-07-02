"""
本体构建器后端 API — 支持文档抽取（extract）模式。
- POST /builder/extract  SSE 流式抽取，从 Word/Excel/PDF/CSV/MD/TXT 文档中
  让 LLM 流式产出 entity_proposed / attr_proposed / relation_proposed
  / rule_suggested / action_suggested 五类事件。
"""
from __future__ import annotations

import json
import logging
import re
from collections.abc import Generator
from io import BytesIO

from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.config import settings
from app.database import get_db
from app.services.copilot import get_llm_client

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/builder", tags=["builder"])


# ── 文档解析 ─────────────────────────────────────────────────────

def _parse_text(filename: str, raw: bytes) -> str:
    """解析单个文件，统一转纯文本喂给 LLM。"""
    name = (filename or "").lower()
    try:
        if name.endswith(".txt") or name.endswith(".md"):
            return raw.decode("utf-8", errors="ignore")

        if name.endswith(".csv"):
            try:
                import pandas as pd  # type: ignore
                df = pd.read_csv(BytesIO(raw))
                return f"CSV: {filename}\n列：{list(df.columns)}\n前 20 行：\n{df.head(20).to_string()}"
            except Exception:
                return raw.decode("utf-8", errors="ignore")

        if name.endswith(".xlsx") or name.endswith(".xls"):
            try:
                import pandas as pd  # type: ignore
                xls = pd.ExcelFile(BytesIO(raw))
                parts = []
                for sheet in xls.sheet_names[:5]:
                    df = xls.parse(sheet)
                    parts.append(f"Sheet: {sheet}\n列：{list(df.columns)}\n前 10 行：\n{df.head(10).to_string()}")
                return f"Excel: {filename}\n\n" + "\n\n".join(parts)
            except Exception as e:
                return f"Excel 解析失败: {e}"

        if name.endswith(".docx"):
            try:
                from docx import Document  # type: ignore
                doc = Document(BytesIO(raw))
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                tables_text = []
                for ti, t in enumerate(doc.tables[:5]):
                    rows = []
                    for row in t.rows[:30]:
                        rows.append(" | ".join(c.text.strip() for c in row.cells))
                    tables_text.append(f"[表 {ti+1}]\n" + "\n".join(rows))
                return f"Word: {filename}\n\n" + "\n".join(paragraphs) + ("\n\n" + "\n\n".join(tables_text) if tables_text else "")
            except Exception as e:
                return f"Word 解析失败: {e}"

        if name.endswith(".pdf"):
            try:
                import pdfplumber  # type: ignore
                with pdfplumber.open(BytesIO(raw)) as pdf:
                    pages = [p.extract_text() or "" for p in pdf.pages[:30]]
                return f"PDF: {filename}\n\n" + "\n\n".join(pages)
            except Exception as e:
                return f"PDF 解析失败: {e}"

        return raw.decode("utf-8", errors="ignore")
    except Exception as e:
        logger.exception("文档解析异常")
        return f"解析失败: {e}"


# ── LLM Prompt ──────────────────────────────────────────────────

_EXTRACT_PROMPT = """你是本体建模专家，需要从用户提供的业务文档中抽取本体草稿。

## 全平台术语规范（严格遵守）
- **对象**（Object）— 业务实体，如客户、订单
- **属性**（Property）— 对象上的字段
- **关系**（Relation）— 对象之间的关联
- **规则**（Rule）— 业务约束，如"客户连续 3 个月不下单视为流失风险"
- **动作**（Action）— 业务操作，如"触发挽留外呼"

## 输入
你将看到一份业务文档的纯文本（可能含有表格、章节）。请通读全文，识别业务对象、属性、关系，
并捕捉文档里隐含的规则与动作建议。

## 输出格式
严格返回 JSON 数组，每个元素是一条事件，type 取以下五种之一：
1. {"type":"entity_proposed","name":"PascalCase英文名","display_name":"中文名","tier":1|2|3,"primary_key":"id","description":"业务描述","icon":"emoji"}
2. {"type":"attr_proposed","entity_name":"必须与上方某个 entity 的 name 一致","name":"snake_case","display_name":"中文","type":"string|number|date|boolean|enum","required":true|false,"description":"说明"}
3. {"type":"relation_proposed","name":"snake_case","display_name":"中文","from_entity":"对象英文名","to_entity":"对象英文名","cardinality":"1:1|1:N|N:N","description":"说明"}
4. {"type":"rule_suggested","name":"中文规则名","description":"业务说明","condition_hint":"触发条件的人话描述","action_hint":"动作描述","target_entity":"建议挂到的对象英文名","source_file":"来源文件名"}
5. {"type":"action_suggested","name":"中文动作名","description":"业务说明","trigger_hint":"何时触发","effect_hint":"产生什么效果","target_entity":"建议挂到的对象英文名","source_file":"来源文件名"}

## 要求
- tier 划分：T1=核心（客户、产品）/ T2=领域（工单、设备）/ T3=场景（分析结果、预警）
- 关系两端必须能在 entity_proposed 里找到同名对象
- 每个对象至少 3 个属性；必有 1 个 required=true 作为主键候选
- **规则、动作只放进 rule_suggested / action_suggested**，不要落入对象的属性
- 不要包含其他文字，只返回 JSON 数组
"""


def _call_llm_for_extract(content: str, filename: str) -> list[dict]:
    client = get_llm_client()
    messages = [
        {"role": "system", "content": _EXTRACT_PROMPT},
        {"role": "user", "content": f"## 文档：{filename}\n\n{content[:18000]}"},
    ]
    try:
        resp = client.chat.completions.create(
            model=settings.LLM_MODEL,
            messages=messages,
            temperature=0,
            max_tokens=6000,
            timeout=90,
        )
    except Exception as e:
        logger.error("LLM 调用失败: %s", e)
        raise
    text = (resp.choices[0].message.content or "").strip()
    # 清理 think 标签和 markdown 代码块
    text = re.sub(r"<think>[\s\S]*?</think>", "", text).strip()
    if text.startswith("```"):
        text = re.sub(r"^```\w*\n", "", text)
        text = re.sub(r"\n```\s*$", "", text)
    # 直接尝试 JSON 数组
    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        # 容错：找首个 [ 与末个 ]
        start = text.find("[")
        end = text.rfind("]")
        if start < 0 or end < 0:
            raise
        data = json.loads(text[start: end + 1])
    if isinstance(data, dict) and "events" in data:
        data = data["events"]
    if not isinstance(data, list):
        return []
    return [e for e in data if isinstance(e, dict) and "type" in e]


# ── SSE 端点 ─────────────────────────────────────────────────────

@router.post("/extract")
async def extract_from_documents(
    session_id: str = Form(""),
    scenario: str = Form(""),
    files: list[UploadFile] = File(default=[]),
    db: Session = Depends(get_db),
):
    """SSE 流式从业务文档抽取本体草稿。"""
    # 先把所有文件读出来（上下文里持久化），避免 yield 时句柄关掉
    docs: list[tuple[str, bytes]] = []
    for f in files:
        if not f.filename:
            continue
        raw = await f.read()
        if raw:
            docs.append((f.filename, raw))

    def _emit(ev: dict) -> str:
        return f"data: {json.dumps(ev, ensure_ascii=False)}\n\n"

    def event_stream() -> Generator[str, None, None]:
        try:
            yield _emit({"type": "extract_started", "session_id": session_id, "file_count": len(docs)})

            if not docs:
                yield _emit({"type": "error", "content": "未收到任何文件"})
                yield "data: [DONE]\n\n"
                return

            # 维护已抛出的对象 name，用于关系/属性合法性校验
            entity_names: set[str] = set()

            for filename, raw in docs:
                yield _emit({"type": "file_parsing", "file_name": filename})
                text = _parse_text(filename, raw)
                if text.startswith("解析失败") or text.startswith("Word 解析失败") or text.startswith("PDF 解析失败"):
                    yield _emit({"type": "file_failed", "file_name": filename, "error": text})
                    continue
                yield _emit({"type": "file_parsed", "file_name": filename, "text_length": len(text)})

                try:
                    events = _call_llm_for_extract(text, filename)
                except Exception as e:
                    yield _emit({"type": "file_failed", "file_name": filename, "error": f"LLM 抽取失败: {e}"})
                    continue

                # 先发对象，再发属性，最后发关系，避免前端拿到孤立属性
                ordered = sorted(events, key=lambda ev: {
                    "entity_proposed": 0,
                    "attr_proposed": 1,
                    "relation_proposed": 2,
                    "rule_suggested": 3,
                    "action_suggested": 4,
                }.get(ev.get("type"), 9))

                for ev in ordered:
                    t = ev.get("type")
                    if t == "entity_proposed":
                        entity_names.add(ev.get("name", ""))
                    elif t == "attr_proposed":
                        if ev.get("entity_name") not in entity_names:
                            continue
                    elif t == "relation_proposed":
                        if ev.get("from_entity") not in entity_names or ev.get("to_entity") not in entity_names:
                            continue
                    elif t in ("rule_suggested", "action_suggested"):
                        ev.setdefault("source_file", filename)
                    yield _emit(ev)

            yield _emit({"type": "extract_finished"})
            yield "data: [DONE]\n\n"

        except Exception as e:
            logger.exception("extract failed")
            yield _emit({"type": "error", "content": str(e)})
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── 对话生成（chat）── 走 AgentService 的 tool calling ────────────

_CHAT_SYSTEM_PROMPT = """你是本体构建器的「场景 Copilot」。你的任务是引导用户从业务场景出发，
联动数据资产与业务文档，生成本体对象、属性、关系、规则建议、动作建议。

## 全平台术语
对象（Object）/ 属性（Property）/ 派生属性（Derived Property）/ 关系（Relation）/ 规则（Rule）/ 动作（Action）。

## 你必须严格按这四阶段推进
1. **business-clarify**：用户描述完业务后，用 1-2 句话总结确认。如果模糊，用一句业务问题澄清（不要提"实体""属性"这些建模术语）。
2. **asset-select**：业务清晰后，调用 `list_business_datasources(domain, keywords)` 得到候选数据源，
   你的回答里只需说一句"为你匹配到了候选数据源，请勾选要纳入本体的资产"。
   不要把数据源名称重复贴出来——前端会渲染 display_card=asset_picker。
3. **doc-select**：用户回答"已选好数据源 [...id列表]"后，调用 `list_business_documents(domain, datasource_ids)` 得到候选文档。
   同样让前端渲染 display_card 让用户勾选。
4. **synthesize**：用户回答"已选好文档 [...id列表]"后，调用
   `analyze_assets_for_ontology(datasource_ids, document_ids, business_context)` 抽取本体草稿。
   返回结果会同时携带规则与动作建议，用户进入审核页时统一处理。

## 几条铁律
- 一次只问一个问题；澄清不超过 2 轮
- 列资产时一定调 list_business_datasources（不要凭印象编名字）
- 列文档时一定调 list_business_documents
- 抽取本体一定调 analyze_assets_for_ontology
- 工具结果里的 display_card 字段保持原样回传，前端会自动渲染
"""


@router.post("/chat")
async def builder_chat(
    session_id: str = Form(""),
    message: str = Form(""),
    db: Session = Depends(get_db),
):
    """对话生成模式 — 用 AgentService 跑 LLM + tool calling，流式 SSE。"""
    from app.services.agent.orchestrator import AgentService

    def _emit(ev: dict) -> str:
        return f"data: {json.dumps(ev, ensure_ascii=False, default=str)}\n\n"

    def event_stream() -> Generator[str, None, None]:
        try:
            agent = AgentService(db, system_prompt_prefix=_CHAT_SYSTEM_PROMPT)
            for ev in agent.ask(message or "（用户没有发送内容，请引导他描述业务场景）"):
                yield _emit(ev)
        except Exception as e:
            logger.exception("builder chat failed")
            yield _emit({"type": "error", "content": str(e)})
        yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── AI 辅助：属性补全 ─────────────────────────────────────────────

class SuggestAttrProperty(BaseModel):
    name: str
    displayName: str | None = None
    type: str = "string"
    required: bool = False


class SuggestAttrRequest(BaseModel):
    object_name: str
    display_name: str | None = None
    tier: int = 2
    description: str | None = None
    existing_properties: list[SuggestAttrProperty] = []


_SUGGEST_ATTR_PROMPT = """你是本体建模专家。给定一个业务对象的基本信息和已有属性，请建议缺失的重要属性。

## 输入
- 对象名（英文 + 中文）
- 对象层级（T1/T2/T3）
- 已有属性列表

## 输出格式
严格返回 JSON 数组，每个元素：
{"name":"snake_case英文名","displayName":"中文名","type":"string|number|date|boolean|enum","required":true|false,"description":"一句话说明"}

## 要求
- 只建议不在已有属性列表中的属性
- 建议 3~8 个最核心的属性
- 属性名用 snake_case，中文名用业务术语
- required=true 只用于业务上必须有值的字段
- 不要包含其他文字，只返回 JSON 数组
"""


@router.post("/suggest-attributes")
async def suggest_attributes(body: SuggestAttrRequest):
    """AI 补全属性：根据对象信息和已有属性，让 LLM 建议缺失属性。"""
    existing_names = {p.name for p in body.existing_properties}
    existing_desc = "\n".join(
        f"  - {p.name} ({p.displayName or p.name}, {p.type}, {'必填' if p.required else '可选'})"
        for p in body.existing_properties
    ) or "  （暂无）"

    user_msg = (
        f"对象：{body.object_name}（{body.display_name or body.object_name}）\n"
        f"层级：T{body.tier}\n"
        f"描述：{body.description or '无'}\n\n"
        f"已有属性：\n{existing_desc}"
    )

    client = get_llm_client()
    try:
        resp = client.chat.completions.create(
            model=settings.LLM_MODEL,
            messages=[
                {"role": "system", "content": _SUGGEST_ATTR_PROMPT},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.1,
            max_tokens=2000,
            timeout=30,
        )
    except Exception as e:
        logger.error("suggest-attributes LLM 调用失败: %s", e)
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=502, content={"detail": f"LLM 调用失败: {e}"})

    text = (resp.choices[0].message.content or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```\w*\n", "", text)
        text = re.sub(r"\n```\s*$", "", text)

    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("[")
        end = text.rfind("]")
        if start < 0 or end < 0:
            return {"suggested_properties": []}
        data = json.loads(text[start:end + 1])

    if not isinstance(data, list):
        data = []

    # 过滤已有属性
    suggested = [
        item for item in data
        if isinstance(item, dict) and item.get("name") and item["name"] not in existing_names
    ]
    return {"suggested_properties": suggested}


# ── AI 辅助：关系推断 ─────────────────────────────────────────────

class SuggestRelObject(BaseModel):
    name: str
    displayName: str | None = None
    tier: int = 2
    properties: list[str] = []


class SuggestRelExisting(BaseModel):
    source: str
    target: str
    name: str | None = None


class SuggestRelRequest(BaseModel):
    objects: list[SuggestRelObject]
    existing_relations: list[SuggestRelExisting] = []


_SUGGEST_REL_PROMPT = """你是本体建模专家。给定一组业务对象及其属性，请推断对象之间可能存在的关系。

## 输入
- 对象列表（含名称、层级、属性名）

## 输出格式
严格返回 JSON 数组，每个元素：
{"name":"snake_case英文名","displayName":"中文名","source":"源对象英文名","target":"目标对象英文名","cardinality":"1:1|1:N|N:N","relationType":"ObjectProperty|SymmetricProperty|TransitiveProperty|FunctionalProperty","semanticType":"composition|event|inheritance|dependency|association","description":"一句话说明"}

## 判断依据
1. 两个对象有同名外键属性（如都有 customer_id）→ 关联关系
2. 层级包含关系（T1 包含 T2）→ 组合关系
3. 业务语义（预警、工单、事件）→ 事件关系
4. 继承关系（子类型）→ 继承关系

## 要求
- source 和 target 必须是输入对象中的 name
- 不要重复已有的关系
- 最多建议 5 条最核心的关系
- 不要包含其他文字，只返回 JSON 数组
"""


@router.post("/suggest-relations")
async def suggest_relations(body: SuggestRelRequest):
    """AI 推断关系：根据对象信息让 LLM 推断可能的关系。"""
    if len(body.objects) < 2:
        return {"suggested_relations": []}

    existing_pairs = {(r.source, r.target) for r in body.existing_relations}

    obj_desc = "\n".join(
        f"  - {o.name}（{o.displayName or o.name}，T{o.tier}）属性：{', '.join(o.properties) or '无'}"
        for o in body.objects
    )
    user_msg = f"对象列表：\n{obj_desc}"

    client = get_llm_client()
    try:
        resp = client.chat.completions.create(
            model=settings.LLM_MODEL,
            messages=[
                {"role": "system", "content": _SUGGEST_REL_PROMPT},
                {"role": "user", "content": user_msg},
            ],
            temperature=0.1,
            max_tokens=2000,
            timeout=30,
        )
    except Exception as e:
        logger.error("suggest-relations LLM 调用失败: %s", e)
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=502, content={"detail": f"LLM 调用失败: {e}"})

    text = (resp.choices[0].message.content or "").strip()
    if text.startswith("```"):
        text = re.sub(r"^```\w*\n", "", text)
        text = re.sub(r"\n```\s*$", "", text)

    try:
        data = json.loads(text)
    except json.JSONDecodeError:
        start = text.find("[")
        end = text.rfind("]")
        if start < 0 or end < 0:
            return {"suggested_relations": []}
        data = json.loads(text[start:end + 1])

    if not isinstance(data, list):
        data = []

    # 过滤已有关系 + 校验 source/target 存在
    obj_names = {o.name for o in body.objects}
    suggested = []
    for item in data:
        if not isinstance(item, dict):
            continue
        src, tgt = item.get("source"), item.get("target")
        if not src or not tgt:
            continue
        if src not in obj_names or tgt not in obj_names:
            continue
        if (src, tgt) in existing_pairs or (tgt, src) in existing_pairs:
            continue
        suggested.append(item)

    return {"suggested_relations": suggested}


# ── AI 辅助：数据源自动映射 ───────────────────────────────────────

class AutoMapProperty(BaseModel):
    name: str
    type: str = "string"
    description: str | None = None


class AutoMapObject(BaseModel):
    id: str
    name: str = ""
    displayName: str = ""
    properties: list[AutoMapProperty] = []


class AutoMapRequest(BaseModel):
    objects: list[AutoMapObject]
    asset_ids: list[str] | None = None
    scope: str | None = None


class _AttrAdapter:
    """轻量 adapter，让 Pydantic 属性能传给 _heuristic_score()。"""
    __slots__ = ("name", "type", "description")
    def __init__(self, name: str, type_: str, desc: str | None):
        self.name = name
        self.type = type_
        self.description = desc or ""


@router.post("/auto-map")
async def auto_map(body: AutoMapRequest, db: Session = Depends(get_db)):
    """智能映射：根据 scope 或 asset_ids 匹配最优表和列。

    三层匹配逻辑：
    1. 确定候选资产范围（scope → connection_id 筛选，或 asset_ids 指定，或全部结构化资产）
    2. 对象名 → 表名 匹配，为每个对象找出最相关的表
    3. 属性 → 列 匹配，在匹配到的表内做列级映射
    """
    from app.models.asset import Asset
    from app.services.data_plane.mapping_suggest_service import _heuristic_score

    if not body.objects:
        return {"mappings": {}}

    # 第一层：确定候选资产范围
    query = db.query(Asset).filter(
        Asset.kind.in_(["table", "sql_view"]),
        Asset.status == "active",
    )
    if body.asset_ids:
        query = query.filter(Asset.id.in_(body.asset_ids))
    elif body.scope:
        query = query.filter(Asset.connection_id == body.scope)

    assets = query.all()
    if not assets:
        return {"mappings": {}}

    # 构建 asset 索引: [(asset_id, table_name, schema)]
    asset_index: list[tuple[str, str, list[dict]]] = []
    for a in assets:
        schema = a.schema_snapshot or []
        if schema:
            asset_index.append((a.id, a.alias or a.name, schema))

    if not asset_index:
        return {"mappings": {}}

    result: dict[str, dict] = {}

    for obj in body.objects:
        obj_label = obj.displayName or obj.name or obj.id

        # 第二层：对象名 → 表名 匹配，找出 top-N 相关表
        table_scores: list[tuple[int, float]] = []
        for idx, (asset_id, table_name, schema) in enumerate(asset_index):
            name_sim = _table_name_similarity(obj_label, obj.name, table_name)
            table_scores.append((idx, name_sim))

        table_scores.sort(key=lambda x: x[1], reverse=True)
        # 取 top 5 候选表（至少 0.2 分以上）
        top_tables = [(idx, score) for idx, score in table_scores[:5] if score >= 0.2]
        if not top_tables:
            # 退化：取前 10 张表全部尝试
            top_tables = table_scores[:10]

        # 第三层：属性 → 列 匹配
        best_total_score = 0.0
        best_result = None
        matched_tables_for_obj: list[dict] = []

        for idx, table_sim in top_tables:
            asset_id, table_name, schema = asset_index[idx]
            obj_score = 0.0
            obj_mappings: dict[str, dict] = {}

            for prop in obj.properties:
                adapter = _AttrAdapter(prop.name, prop.type, prop.description)
                best_s = 0.0
                best_col = None
                for col in schema:
                    s, _r = _heuristic_score(adapter, col)
                    if s > best_s:
                        best_s = s
                        best_col = col.get("name")
                if best_col and best_s > 0:
                    obj_mappings[prop.name] = {
                        "column": best_col,
                        "score": round(best_s, 3),
                        "asset_id": asset_id,
                    }
                    obj_score += best_s

            combined_score = obj_score + table_sim * len(obj.properties) * 0.3
            matched_tables_for_obj.append({
                "asset_id": asset_id,
                "table_name": table_name,
                "columns": [col.get("name", "") for col in schema],
                "score": round(combined_score, 3),
            })

            if combined_score > best_total_score:
                best_total_score = combined_score
                best_result = {"asset_id": asset_id, "mappings": obj_mappings}

        # 按分数排序，返回 top 相关表
        matched_tables_for_obj.sort(key=lambda x: x["score"], reverse=True)
        top_matched = matched_tables_for_obj[:5]

        if best_result:
            result[obj.id] = {
                "asset_id": best_result["asset_id"],
                "matched_tables": top_matched,
                "property_mappings": best_result["mappings"],
            }
        else:
            result[obj.id] = {
                "asset_id": None,
                "matched_tables": top_matched,
                "property_mappings": {},
            }

    return {"mappings": result}


def _table_name_similarity(obj_display: str, obj_name: str, table_name: str) -> float:
    """计算对象名与表名的相似度。综合中文名匹配和英文名匹配。"""
    import re
    from difflib import SequenceMatcher

    table_lower = table_name.lower().replace("_", "").replace("-", "")
    scores: list[float] = []

    if obj_name:
        obj_lower = obj_name.lower().replace("_", "").replace("-", "")
        if obj_lower == table_lower:
            return 1.0
        if obj_lower in table_lower or table_lower in obj_lower:
            scores.append(0.8)
        scores.append(SequenceMatcher(None, obj_lower, table_lower).ratio())

    if obj_display:
        display_lower = obj_display.lower()
        if display_lower in table_lower or table_lower in display_lower:
            scores.append(0.7)
        # 拆分 token 比较
        obj_tokens = set(re.findall(r'[一-鿿]+|[a-z0-9]+', display_lower))
        table_tokens = set(re.findall(r'[一-鿿]+|[a-z0-9]+', table_lower))
        if obj_tokens and table_tokens:
            overlap = len(obj_tokens & table_tokens)
            if overlap:
                scores.append(overlap / max(len(obj_tokens), len(table_tokens)) * 0.75)

    return max(scores) if scores else 0.0


# ── AI 辅助：字段推荐 ────────────────────────────────────────────────

class SuggestColumnsRequest(BaseModel):
    asset_id: str
    properties: list[dict]  # [{name, type?, description?}]


@router.post("/suggest-columns")
def suggest_columns(body: SuggestColumnsRequest, db: Session = Depends(get_db)):
    """给定属性列表 + 资产 ID，返回每属性的字段推荐排序。"""
    from app.services.data_plane.asset_service import AssetService
    from app.services.data_plane.mapping_suggest_service import _heuristic_score

    asset = AssetService(db).get(body.asset_id)
    if not asset:
        raise HTTPException(404, "资产不存在")

    columns = asset.schema_snapshot or []
    results = []

    for prop in body.properties:

        class _FakeAttr:
            pass

        attr = _FakeAttr()
        attr.name = prop.get("name", "")
        attr.type = prop.get("type", "")
        attr.description = prop.get("description", "")

        candidates = []
        for col in columns:
            score, reason = _heuristic_score(attr, col)
            if score > 0.1:
                tier = "high" if score >= 0.7 else "medium" if score >= 0.4 else "low"
                candidates.append({
                    "column": col.get("name", ""),
                    "column_type": col.get("type"),
                    "comment": col.get("comment"),
                    "score": round(score, 3),
                    "tier": tier,
                    "reason": reason,
                })

        candidates.sort(key=lambda c: c["score"], reverse=True)
        results.append({
            "property_name": prop.get("name", ""),
            "candidates": candidates[:5],
        })

    return {"suggestions": results}


# ── 水合演练（hydrate）── 验证本体草稿与真实数据的映射 ──────────────

@router.post("/hydrate")
async def builder_hydrate(body: dict, db: Session = Depends(get_db)):
    """SSE 流式水合演练：验证本体草稿与真实数据的映射。"""
    from app.services.builder.hydration_service import HydrationRequest, HydrationService

    try:
        req = HydrationRequest(**body)
    except Exception as e:
        from fastapi.responses import JSONResponse
        return JSONResponse(status_code=422, content={"detail": str(e)})

    def _emit(ev: dict) -> str:
        return f"data: {json.dumps(ev, ensure_ascii=False, default=str)}\n\n"

    def event_stream() -> Generator[str, None, None]:
        try:
            yield _emit({"type": "hydrate_started", "session_id": req.session_id})
            svc = HydrationService(db)
            for ev in svc.run(req):
                yield _emit(ev)
            yield "data: [DONE]\n\n"
        except Exception as e:
            logger.exception("hydrate failed")
            yield _emit({"type": "error", "content": str(e)})
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"},
    )


# ── M3: 发布草稿 → OntologyEntity + EntityAttribute + ObjectBinding ─────

class _DraftProperty(BaseModel):
    name: str
    displayName: str | None = None
    type: str = "string"
    required: bool = False
    description: str | None = None
    source_asset_id: str | None = None
    source_column: str | None = None


class _DraftRelation(BaseModel):
    name: str
    displayName: str | None = None
    source: str
    target: str
    cardinality: str = "1:N"
    relationType: str = "ObjectProperty"
    description: str | None = None


class _DraftObject(BaseModel):
    name: str
    displayName: str | None = None
    tier: int = 2
    namespace: str | None = None
    description: str | None = None
    primaryKey: str | None = None
    properties: list[_DraftProperty] = []
    backing_asset_ids: list[str] = []


class FinalizeRequest(BaseModel):
    objects: list[_DraftObject]
    relations: list[_DraftRelation] = []
    # 本批发布对象统一归属的场景 code 列表（与已存在场景做并集写入）
    scenario_codes: list[str] = []


@router.post("/finalize")
def builder_finalize(body: FinalizeRequest, db: Session = Depends(get_db)):
    """把 builder 草稿一次性落库：每个对象创建 OntologyEntity + EntityAttribute；
    若 backing_asset_ids 非空，自动创建 primary ObjectBinding（含 field_mappings）。

    幂等：实体已存在时跳过；binding 已存在时也跳过。
    """
    from app.models.entity import EntityAttribute, OntologyEntity
    from app.repositories.asset_repo import AssetRepository
    from app.services.data_plane.object_binding_service import ObjectBindingService

    asset_repo = AssetRepository(db)
    binding_svc = ObjectBindingService(db)

    # 本批发布统一归属的场景（去重、去空）
    batch_scenarios = [c for c in dict.fromkeys(body.scenario_codes) if c]

    created_entities = 0
    created_bindings = 0
    skipped = 0

    for obj in body.objects:
        eid = (f"{obj.namespace}_{obj.name}" if obj.namespace else obj.name)
        existing = db.get(OntologyEntity, eid)
        if existing:
            skipped += 1
            entity = existing
        else:
            entity = OntologyEntity(
                id=eid, name=obj.name, name_cn=obj.displayName or obj.name,
                tier=int(obj.tier or 2), status="active",
                description=obj.description,
                config_json={"primary_key": obj.primaryKey,
                             "backing_asset_ids": obj.backing_asset_ids},
            )
            db.add(entity)
            db.flush()
            created_entities += 1

        # 场景归属：与对象已有 scenario_codes 做并集，保持顺序
        if batch_scenarios:
            merged = list(dict.fromkeys([*(entity.scenario_codes or []), *batch_scenarios]))
            entity.scenario_codes = merged
            db.flush()

        # 属性（按 name 去重）
        attr_by_name: dict[str, EntityAttribute] = {a.name: a for a in entity.attributes}
        for p in obj.properties:
            if p.name in attr_by_name:
                continue
            attr = EntityAttribute(
                entity_id=entity.id, name=p.name, type=p.type,
                description=p.description or p.displayName or p.name,
                required=p.required,
            )
            db.add(attr)
            db.flush()
            attr_by_name[p.name] = attr
        db.flush()

        # ObjectBinding（每个 backing_asset_id 一个 primary，role 区分）
        for idx, asset_id in enumerate(obj.backing_asset_ids):
            asset = asset_repo.get_by_id(asset_id)
            if not asset:
                continue
            role = "primary" if idx == 0 else "enrichment"
            existing_binding = binding_svc.repo.find_existing(entity.id, asset_id, role)
            if existing_binding:
                continue
            field_mappings = []
            for p in obj.properties:
                if p.source_asset_id == asset_id and p.source_column:
                    attr = attr_by_name.get(p.name)
                    if attr:
                        field_mappings.append({
                            "attribute_id": attr.id,
                            "source_column": p.source_column,
                            "transform": None,
                        })
            try:
                binding_svc.create(
                    object_type_id=entity.id,
                    asset_id=asset_id,
                    role=role,
                    field_mappings=field_mappings,
                    id_column=obj.primaryKey,
                )
                created_bindings += 1
            except Exception:
                logger.exception("create binding failed for %s", entity.id)

    # ── 关系持久化 ──
    from app.models.relation import EntityRelation

    created_relations = 0
    entity_id_map = {}
    for obj in body.objects:
        eid = (f"{obj.namespace}_{obj.name}" if obj.namespace else obj.name)
        entity_id_map[obj.name] = eid

    for rel in body.relations:
        from_eid = entity_id_map.get(rel.source)
        to_eid = entity_id_map.get(rel.target)
        if not from_eid or not to_eid:
            continue
        existing_rel = db.query(EntityRelation).filter(
            EntityRelation.from_entity_id == from_eid,
            EntityRelation.to_entity_id == to_eid,
            EntityRelation.name == rel.name,
        ).first()
        if existing_rel:
            continue
        new_rel = EntityRelation(
            from_entity_id=from_eid,
            to_entity_id=to_eid,
            name=rel.name,
            rel_type=rel.relationType or "ObjectProperty",
            cardinality=rel.cardinality or "1:N",
            description=rel.description or rel.displayName or rel.name,
        )
        db.add(new_rel)
        created_relations += 1

    db.commit()
    return {
        "created_entities": created_entities,
        "created_bindings": created_bindings,
        "created_relations": created_relations,
        "skipped": skipped,
        "total": len(body.objects),
    }
